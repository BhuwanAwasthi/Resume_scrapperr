import os
import re
import pdfplumber
from docx import Document
from flask import Flask, request, send_from_directory, render_template, abort
import pandas as pd

app = Flask(__name__)

# Define folders and ensure they exist
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

def extract_text_from_pdf(pdf_path):
    all_text = ''
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            all_text += page.extract_text() + ' '
    return all_text

def extract_text_from_docx(docx_path):
    all_text = ''
    doc = Document(docx_path)
    for para in doc.paragraphs:
        all_text += para.text + '\n'
    for section in doc.sections:
        for header in section.header.paragraphs:
            all_text += header.text + '\n'
        for footer in section.footer.paragraphs:
            all_text += footer.text + '\n'
    return all_text

def extract_text_from_file(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")

def find_emails(text):
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    emails = re.findall(email_pattern, text)
    return list(set(emails))

def clean_email(emails):
    cleaned_emails = []
    for email in emails:
        if '-' in email:
            cleaned_email = email.split('-')[-1]
        else:
            cleaned_email = email
        if re.match(r'[\w\.-]+@[\w\.-]+\.\w+', cleaned_email):
            cleaned_emails.append(cleaned_email)
    return cleaned_emails

def find_phone_numbers(text):
    phone_pattern = r'\+?\d{1,3}[ -]?[\(]?\d{1,3}[\)]?[ -]?[\d -]{7,10}\d'
    phone_numbers = re.findall(phone_pattern, text)
    cleaned_phone_numbers = [''.join(filter(str.isdigit, num)) for num in phone_numbers]
    return list(set(cleaned_phone_numbers))

def filter_valid_phone_numbers(phone_numbers):
    valid_pattern = r'^(91\d{10})$|^[987]\d{9}$'
    valid_phones = [num for num in phone_numbers if re.match(valid_pattern, num)]
    return valid_phones

def extract_information_from_file(file_path):
    text = extract_text_from_file(file_path)
    raw_emails = find_emails(text)
    cleaned_emails = clean_email(raw_emails)
    phone_numbers = find_phone_numbers(text)
    filtered_phones = filter_valid_phone_numbers(phone_numbers)
    return {"Emails": cleaned_emails, "Phone Numbers": filtered_phones, "Text": text}

@app.route('/')
def upload_page():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    files = request.files.getlist('file')
    results = {}
    for file in files:
        if file:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            result = extract_information_from_file(file_path)
            results[file.filename] = result
            os.remove(file_path)  # Optionally remove file after processing
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], 'output.xlsx')
    data_to_excel(results, output_path)
    if not os.path.exists(output_path):
        abort(404, description="File not found.")
    return send_from_directory(app.config['OUTPUT_FOLDER'], 'output.xlsx', as_attachment=True)

def data_to_excel(data, output_path):
    if os.path.exists(output_path):
        os.remove(output_path)  # Remove the old output file if it exists
    
    # Prepare data for DataFrame
    rows = []
    for filename, info in data.items():
        row = {
            "File Name": filename,
            "Email": ', '.join(info['Emails']),
            "Phone": ', '.join(info['Phone Numbers']),
            "Resume Text": info['Text']
        }
        rows.append(row)
    
    # Create DataFrame
    df = pd.DataFrame(rows)
    
    # Write DataFrame to Excel
    with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Extracted Data', index=False)


@app.route('/clear', methods=['GET'])
def clear_data():
    # Clear uploaded files
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        os.remove(file_path)
    # Clear output files
    for filename in os.listdir(app.config['OUTPUT_FOLDER']):
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        os.remove(file_path)
    return "Data cleared successfully", 200

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))  # Get port from environment variable or default to 5000
    app.run(host='0.0.0.0', port=port, debug=True)  # Bind to all IP addresses, use the correct port
